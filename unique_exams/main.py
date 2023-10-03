from docx import Document
import random
import os
from typing import List
from config import ALTERNATIVE_TEXT, END_ALTERNATIVES
from token_hash import to_hash, from_hash

def load_in_template(filename)->dict:
    """ loads in the content of the file to memory, takes note of the choices and where they begin and end, returns a dict with all the loaded content """
    # Output format of dict `loaded_content`
    # loaded_content['common'] : contains common paragraphs as a list.
    # loaded_content['choices'] : a list of choice dicts
    # choice_dict format:
    # { 'after_paragraph':int # the paragraph index of 'common' which this choice comes after 
    # 'alternatives': list # list of alternative content in format alternative_content
    # }
    # alternative_content format: A list of paragraphs

    template_doc = Document(filename)
    loaded_content = {'common':[], 'choices':[]}

    print(f"==Loading in {filename}==")

    mode = 'common'
    for i, paragraph in enumerate(template_doc.paragraphs):
        print(f"{i}:{mode}:       {paragraph.text[:45]}")
        if ALTERNATIVE_TEXT.lower() in paragraph.text.lower():
            if mode == 'common':
                # If we encounter a choice header while in common mode
                alternatives = [] # start a new alternatives list
                paragraph_index = len(loaded_content['common'])
                print(f"Detected Start of choices")
            else:
                # we save the current alternative
                alternatives.append(alternative_content)
                print(f"Detected next alternative choice")
            mode = 'choice'
            alternative_content = [] # we are starting a new alternative
            continue # we don't save the header
        if mode == 'choice':
            if END_ALTERNATIVES.lower() in paragraph.text.lower():
                # save current alternative and exit choice mode
                alternatives.append(alternative_content)
                # exit choice mode
                loaded_content['choices'].append({
                    'after_paragraph': paragraph_index,
                    'alternatives': alternatives
                })
                mode = 'common'
                print(f"Detected end of choices")
            else:
                # add paragraph to current choice
                alternative_content.append(paragraph)
                print(f"Alternative text added")
        else:
            # add paragraph to common
            loaded_content['common'].append(paragraph)
            print(f"Common text added")
    for choice in loaded_content['choices']:
        print (f"After paragraph {choice['after_paragraph']} there are {len(choice['alternatives'])} alternatives")
    return loaded_content

def make_choices(content)->List[int]:
    """ takes the loaded content dict and randomly assigns choices for each section"""
    choices = []
    for choice in content['choices']:
        random_index = random.randint(0, len(choice['alternatives']) - 1)
        choices.append(random_index)
    return choices

def create_new_document(content:dict, choices:List[int], filename:str, original:str):
    file_path = os.path.join(os.getcwd(), "output", f"{filename}")
    # we begin by deep copying the original document to move styles across
    original_document = Document(original)
    document = Document()

    # The following are currently disabled as they are presently not working
    # document.settings = original_document.settings
    # document.part = original_document.part
    # document.core_properties = original_document.core_properties
    # add_tables(document, original_document) # Currently disabled, as it seems to add the tables to the top of document, not in their correct locations
    # for section in original_document.sections:
    #     document.add_section(section.start_type, section.continuous_section_break, section.orientation)     
    for shaoe in original_document.inline_shapes:
        document.add_picture(shaoe._inline.graphic.graphicData.pic.nvPicPr.cNvPr.get('name'), width=shaoe.width, height=shaoe.height)

    document.save(file_path)

    # now we add the new content
    document = Document(file_path)
    choice_content = {}
    for i, choice in enumerate(content['choices']):
        after_paragraph = choice['after_paragraph']
        alternative_index = choices[i]

        try:
            alternative_content = choice['alternatives'][alternative_index]
        except:
            raise Exception(f"Alternative {i} does not exist for {len(choice['alternatives'])} alternatives")
        choice_content[after_paragraph] = alternative_content

    for i, paragraph in enumerate(content['common']):
        # we check if there is a choice to be inserted before adding this paragraph
        if i in choice_content.keys():
            # get the choice content
            content = choice_content[i]
            # add choice content instead of the next paragraph
            for choice_paragraph in content:
                add_paragraph(document, choice_paragraph)                
        # add the common paragraph
        add_paragraph(document, paragraph)
    document.save(file_path)
    print(f"Saved to {file_path}")

def add_tables(target_document, original_document):
    """ Adds tables from the original document to the target document"""
    
    for table in original_document.tables:
        row_len = len(table.rows)
        col_len = len(table.columns)
        new_table = target_document.add_table(rows=row_len, cols=col_len, style=table.style)
        for r in range(row_len):
            for c in range(col_len):
                try:
                    original_cell_paragraph = table._rows[r].cells[c].paragraphs[0]
                except IndexError:
                    # cell is empty
                    continue
                new_table.rows[r].cells[c].add_paragraph(text=original_cell_paragraph.text, style=original_cell_paragraph.style)

def add_style(target_document, target_paragraph, source_style):
    # TODO: This doesn't seem to be working ...
    try:
        target_style = target_document.styles[source_style.name]
    except KeyError:
        # style does not exist, create it
        new_style = target_document.styles.add_style(source_style.name, source_style.type)
        # add attributes
        new_style.font.name = style.font.name
        new_style.font.size = style.font.size
        new_style.font.bold = style.font.bold
        new_style.font.italic = style.font.italic
        new_style.font.underline = style.font.underline
        new_style.font.color.rgb = style.font.color.rgb
        new_style.paragraph_format.alignment = style.paragraph_format.alignment
        new_style.paragraph_format.first_line_indent = style.paragraph_format.first_line_indent
        new_style.paragraph_format.line_spacing = style.paragraph_format.line_spacing
        new_style.paragraph_format.line_spacing_rule = style.paragraph_format.line_spacing_rule
        new_style.paragraph_format.left_indent = style.paragraph_format.left_indent
        new_style.paragraph_format.right_indent = style.paragraph_format.right_indent
        new_style.paragraph_format.space_after = style.paragraph_format.space_after
        new_style.paragraph_format.space_before = style.paragraph_format.space_before
        new_style.paragraph_format.widow_control = style.paragraph_format.widow_control
        new_style.base_style = style.base_style
        new_style.hidden = style.hidden
        new_style.locked = style.locked
        new_style.quick_style = style.quick_style
        new_style.semi_hidden = style.semi_hidden
        new_style.style_id = style.style_id
        new_style.style_type = style.style_type
        new_style.next_style = style.next_style
        new_style.parent_style = style.parent_style
    target_paragraph.style = target_document.styles[source_style.name]    

def add_paragraph(target_document, source_paragraph):
    """ Attempt to copy all formatting over from the original paragraph to a new paragraph appended at the end of the target document """
    print(f"Adding: {source_paragraph.text:40}")
    new_paragraph = target_document.add_paragraph()
    add_style(target_document, new_paragraph, source_paragraph.style)
    # copy runs across (text and formatting as it varies throughout the paragraph)
    for run in source_paragraph.runs:
        new_run = new_paragraph.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.name = run.font.name
        new_run.font.size = run.font.size
        new_run.font.color.rgb = run.font.color.rgb
    # copy paragraph formatting across
    format = new_paragraph.paragraph_format
    source_f = source_paragraph.paragraph_format
    for attribute in ["first_line_indent", "keep_together", "keep_with_next", "left_indent", 
                    "line_spacing", "line_spacing_rule", "page_break_before", "right_indent", 
                    "space_after", "space_before", "widow_control"]:
        setattr(format, attribute, getattr(source_f, attribute))
    return new_paragraph

if __name__ == '__main__':
    # new_exams = int(input("How many new exams do you want to generate? "))
    new_exams = 3
    print(os.getcwd())
    if "unique_exams" in os.getcwd():
        INPUT_FOLDER = os.path.join(os.getcwd(), "input")
        OUTPUT_FOLDER = os.path.join(os.getcwd(), "output")
    else:
        INPUT_FOLDER = os.path.join(os.getcwd(), "unique_exams", "input")
        OUTPUT_FOLDER = os.path.join(os.getcwd(), "unique_exams", "output")
    
    exam_names = [f for f in os.listdir(INPUT_FOLDER) if f.endswith(".docx")]    
    exam_paths = [os.path.join(INPUT_FOLDER, f) for f in exam_names]
    # remove file endings
    exam_names = [os.path.splitext(f)[0] for f in exam_names]
    loaded_contents = [load_in_template(os.path.join(INPUT_FOLDER, f)) for f in exam_paths]
    exam_structures = []
    for loaded_content in loaded_contents:
        exam_structures.append([len(choice['alternatives']) for choice in loaded_content['choices']])
    for i in range(len(exam_structures)):
        if exam_structures[i] != exam_structures[0]:
            raise Exception("The structure of the exam versions are not the same. Please check the files")
    for i in range(new_exams):
        choices = make_choices(loaded_contents[0])
        hash = to_hash(choices)
        reversed_from_hash = from_hash(hash)
        print (f"Choices: {choices} -> {hash} -> {reversed_from_hash}")
        for j, loaded_content in enumerate(loaded_contents):
            exam_name = exam_names[j]
            output_path = os.path.join(OUTPUT_FOLDER, f"{exam_name} - version {hash}.docx")
            create_new_document(loaded_content, choices, output_path, original=exam_paths[j])
