from docx import Document
import random
import os
from typing import List

def load_in_template(filename)->dict:
    """ loads in the content of the file to memory, takes note of the choices and where they begin and end, returns a dict with all the loaded content """
    # Output format of dict `loaded_content`
    # loaded_content['common'] : contains common paragraphs as a list.
    # loaded_content['choices'] : a list of choice dicts
    # choice_dict format:
    # { 'after_paragraph':int # the paragraph index of 'common' which this choice comes after 
    # 'alternatives': list # list of alternative content in format alternative_content
    # }
    # alternative_content format: A list of paragraphs

    template_doc = Document(filename)
    loaded_content = {'common':[], 'choices':[]}

    print(f"==Loading in {filename}==")

    mode = 'common'
    for i, paragraph in enumerate(template_doc.paragraphs):
        # if paragraph text contains `Alternatif Soru`...
        print(f"{i}:{mode}:       {paragraph.text[:45]}")
        if 'alternatif soru' in paragraph.text.lower():
            if mode == 'common':
                # If we encounter a choice header while in common mode
                alternatives = [] # start a new alternatives list
                paragraph_index = len(loaded_content['common'])
                print(f"Detected Start of choices")
            else:
                # we save the current alternative
                alternatives.append(alternative_content)
                print(f"Detected next Alternatif Soru")
            mode = 'choice'
            alternative_content = [] # we are starting a new alternative
            continue # we don't save the header
        if mode == 'choice':
            if 'alternatif sonu' in paragraph.text.lower():
                # save current alternative and exit choice mode
                alternatives.append(alternative_content)
                # exit choice mode
                loaded_content['choices'].append({
                    'after_paragraph': paragraph_index,
                    'alternatives': alternatives
                })
                mode = 'common'
                print(f"Detected end of choices")
            else:
                # add paragraph to current choice
                alternative_content.append(paragraph)
                print(f"Alternative text added")
        else:
            # add paragraph to common
            loaded_content['common'].append(paragraph)
            print(f"Common text added")
    for choice in loaded_content['choices']:
        print (f"After paragraph {choice['after_paragraph']} there are {len(choice['alternatives'])} alternatives")
    return loaded_content

def make_choices(content)->List[int]:
    """ takes the loaded content dict and randomly assigns choices for each section"""
    choices = []
    for choice in content['choices']:
        random_index = random.randint(0, len(choice['alternatives']) - 1)
        choices.append(random_index)
    return choices

def create_new_document(content:dict, choices:List[int], filename:str, original:str):
    file_path = os.path.join(os.getcwd(), "output", f"{filename}")
    # we begin by deep copying the original document to move styles across
    original_document = Document(original)
    document = Document()

    # copy settings and styles across
    for style in original_document.styles:
        try:
	        document.styles.add_style(style.name, style.type)
        except ValueError:
            # style already exists
            pass
    # document.settings = original_document.settings
    # document.part = original_document.part
    # document.core_properties = original_document.core_properties
    for table in original_document.tables:
        row_len = len(table.rows)
        col_len = len(table.columns)
        new_table = document.add_table(rows=row_len, cols=col_len, style=table.style)
        for r in range(row_len):
            for c in range(col_len):
                try:
                    original_cell_paragraph = table._rows[r].cells[c].paragraphs[0]
                except IndexError:
                    # cell is empty
                    continue
                new_table.rows[r].cells[c].add_paragraph(text=original_cell_paragraph.text, style=original_cell_paragraph.style)
    for shaoe in original_document.inline_shapes:
        document.add_picture(shaoe._inline.graphic.graphicData.pic.nvPicPr.cNvPr.get('name'), width=shaoe.width, height=shaoe.height)
    # for section in original_document.sections:
    #     document.add_section(section.start_type, section.continuous_section_break, section.orientation)     
    document.save(file_path)

    # now we add the new content
    document = Document(file_path)
    choice_content = {}
    for i, choice in enumerate(content['choices']):
        after_paragraph = choice['after_paragraph']
        alternative_index = choices[i]

        try:
            alternative_content = choice['alternatives'][alternative_index]
        except:
            raise Exception(f"Alternative {i} does not exist for {len(choice['alternatives'])} alternatives")
        choice_content[after_paragraph] = alternative_content

    for i, paragraph in enumerate(content['common']):
        # we check if there is a choice to be inserted before adding this paragraph
        if i in choice_content.keys():
            # get the choice content
            content = choice_content[i]
            # add choice content instead of the next paragraph
            for choice_paragraph in content:
                add_paragraph(document, choice_paragraph)                
        # add the common paragraph
        add_paragraph(document, paragraph)
    document.save(file_path)
    print(f"Saved to {file_path}")

def add_paragraph(target_document, source_paragraph):
    # method A
    # document.add_paragraph(choice_paragraph.text, style=choice_paragraph.style)

    # method B
    new_paragraph = target_document.add_paragraph()
    new_paragraph.style = source_paragraph.style
    for run in source_paragraph.runs:
        new_run = new_paragraph.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.name = run.font.name
        new_run.font.size = run.font.size
        new_run.font.color.rgb = run.font.color.rgb
    return new_paragraph

import string
HEX_ENCODING = "tpbacdomuginesAEIOUMDR"

def to_hash(choices):
    hex_choices = [HEX_ENCODING[num] for num in choices] 
    return "".join(hex_choices)
    

def from_hash(hash):
    return [HEX_ENCODING.index(char) for char in hash]


if __name__ == '__main__':
    # new_exams = int(input("How many new exams do you want to generate? "))
    new_exams = 3
    print(os.getcwd())
    if "exam_creator" in os.getcwd():
        INPUT_FOLDER = os.path.join(os.getcwd(), "input")
        OUTPUT_FOLDER = os.path.join(os.getcwd(), "output")
    else:
        # add to path exam_creator
        INPUT_FOLDER = os.path.join(os.getcwd(), "exam_creator", "input")
        OUTPUT_FOLDER = os.path.join(os.getcwd(), "exam_creator", "output")
    examiner_path = os.path.join(INPUT_FOLDER, "JRSPA - Examiner Version - simplified.docx")
    student_path = os.path.join(INPUT_FOLDER, "JRSPA - Student Version.docx")
    examiner_loaded_content = load_in_template(examiner_path)
    student_loaded_content = load_in_template(student_path)

    examiner_structure = [len(choice['alternatives']) for choice in examiner_loaded_content['choices']]
    student_structure = [len(choice['alternatives']) for choice in student_loaded_content['choices']]
    if examiner_structure != student_structure:
        raise Exception("The structure of the examiner and student versions are not the same. Please check the files")

    for i in range(new_exams):
        choices = make_choices(student_loaded_content)
        hash = to_hash(choices)
        reversed_from_hash = from_hash(hash)
        print (f"Choices: {choices} -> {hash} -> {reversed_from_hash}")
        # choices is an array like [3,0,2,5], we now convert this into a reversable hash
        examiner_output_path = os.path.join(OUTPUT_FOLDER, f"Examiner version {hash}.docx")
        student_output_path = os.path.join(OUTPUT_FOLDER, f"Student version {hash}.docx")
        create_new_document(student_loaded_content, choices, student_output_path, original=student_path)
        create_new_document(examiner_loaded_content, choices, examiner_output_path, original=examiner_path)
